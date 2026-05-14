from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from bot.utils.logger import logger


def export_docx(text: str, output_path: Path, filename: str = "") -> Path:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)

    if filename:
        heading = doc.add_heading(filename, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        lines = para_text.strip().split("\n")
        table_lines = _collect_table_lines(lines)
        if table_lines:
            _add_table_to_doc(doc, table_lines)
            continue
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            else:
                p = doc.add_paragraph(line)
                p.space_after = Pt(6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.debug("DOCX exportado: {}", output_path)
    return output_path


def _collect_table_lines(lines: list[str]) -> list[list[str]] | None:
    table_data: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            if all(c in "| -:" for c in stripped) and "---" in stripped:
                continue
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            if len(cells) >= 2:
                table_data.append(cells)
    return table_data if table_data else None


def _add_table_to_doc(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()
