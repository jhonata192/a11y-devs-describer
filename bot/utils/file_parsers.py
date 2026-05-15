from pathlib import Path
from bot.utils.logger import logger


def extract_text_from_docx(file_path: Path) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    paragraphs.append("| " + " | ".join(cells) + " |")

        text = "\n\n".join(paragraphs)
        logger.info("Texto extraido de DOCX: {} chars", len(text))
        return text
    except ImportError:
        logger.error("python-docx nao disponivel para extracao")
        return ""


def extract_text_from_html(file_path: Path) -> str:
    try:
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self._text = []
                self._skip = False

            def handle_starttag(self, tag, attrs):
                if tag in ("script", "style", "nav"):
                    self._skip = True

            def handle_endtag(self, tag):
                if tag in ("script", "style", "nav"):
                    self._skip = False
                if tag in ("p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "br", "tr"):
                    self._text.append("\n")

            def handle_data(self, data):
                if not self._skip:
                    stripped = data.strip()
                    if stripped:
                        self._text.append(stripped + " ")

            def get_text(self):
                return "\n".join(
                    line.strip()
                    for line in "".join(self._text).split("\n")
                    if line.strip()
                )

        content = file_path.read_text(encoding="utf-8", errors="replace")
        parser = TextExtractor()
        parser.feed(content)
        text = parser.get_text()
        logger.info("Texto extraido de HTML: {} chars", len(text))
        return text
    except ImportError:
        logger.error("HTML parser nao disponivel")
        return ""


def extract_text_from_file(file_path: Path) -> str | None:
    ext = file_path.suffix.lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".html":
        return extract_text_from_html(file_path)
    return None
