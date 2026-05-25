from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from pipeline.verbosity_manager import filter_blocks_for_profile


def render_docx(document: dict[str, Any], output_path: Path, profile_name: str = "docx", filename: str = "") -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    if filename:
        heading = doc.add_heading(filename, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for section in document.get("sections", []):
        _render_section(doc, section, profile_name)
    doc.save(str(output_path))
    return output_path


def _render_section(doc: Document, section: dict[str, Any], profile_name: str) -> None:
    if section.get("title"):
        doc.add_heading(section["title"], level=min(section.get("level", 1), 9))
    for block in filter_blocks_for_profile(section.get("blocks", []), profile_name):
        _render_block(doc, block)
    for child in section.get("children", []):
        _render_section(doc, child, profile_name)


def _render_block(doc: Document, block: dict[str, Any]) -> None:
    block_type = block.get("type")
    if block_type == "heading":
        doc.add_heading(block.get("title", block.get("text", "")), level=min(block.get("level", 1), 9))
    elif block_type == "paragraph":
        doc.add_paragraph(block.get("text", ""))
    elif block_type == "code":
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(block.get("text", ""))
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        paragraph.style = doc.styles["No Spacing"]
    elif block_type == "list":
        style = "List Number" if block.get("ordered") else "List Bullet"
        for item in block.get("items", []):
            doc.add_paragraph(str(item), style=style)
    elif block_type == "table":
        rows = block.get("rows", [])
        if rows:
            table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = str(cell)
    elif block_type in {"details", "note", "warning", "quote", "image", "math"}:
        doc.add_paragraph(block.get("text", block.get("alt_text", "")))
    else:
        doc.add_paragraph(block.get("text", ""))